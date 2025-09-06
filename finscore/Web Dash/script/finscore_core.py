# %% [markdown]
# <a href='https://colab.research.google.com/github/fernandovieira1/FinScore/blob/main/FINSCORE.ipynb' target='_parent'><img src='https://colab.research.google.com/assets/colab-badge.svg' alt='Open In Colab'/></a>

# %%
# Baseado em 'FINSCORE_v9.6.ipynb'

# %% [markdown]
# **# INTRUÇÕES**

# %% [markdown]
# * Antes de iniciar, certifique-se de estar logado na sua conta Google.
# 
# * Um botão azul 'Fazer login', localizado no canto superior direito da tela, aparecerá, se não.
# 
# * Feito isto, basta inserir os dados nas seções abaixo descritas:
#     - 1.1 Cliente
#     - 1.2 Período
#     - 1.3 Score Serasa
#     - 1.4 Lançamento dos dados Contábeis (na planilha indicada no link)
# 
# * Cada uma das seções mencionadas possui anotações sobre como proceder.
# 
# * Logo após, clique no menu 'Ambiente de execução' e em 'Executar tudo' (ou CTR + F9), nesta ordem.

# %% [markdown]
# #### 0. LANÇAMENTO DOS DADOS #####

# %% [markdown]
# ##### 0.1 Cliente

# %% [markdown]
# *--> Insira o nome do cliente/empresa*

# %%
# Cliente
cliente = 'CARGOBR TRANSPORTES'

# %% [markdown]
# ##### 0.2 Período

# %% [markdown]
# *--> Insira os anos inicial (a primeira) e final (da última) das demonstrações contábeis que serviram de base para a análise*

# %%
# Ano Inicial
ano_inicial = 2021

# %%
# Ano Final
ano_final = 2023

# %% [markdown]
# ##### 0.3 Score Serasa

# %% [markdown]
# *--> Insira o score Serasa do cliente*

# %%
# Serasa
serasa = 550

# %% [markdown]
# ##### 0.4 Lançamento dos dados Contábeis

# %% [markdown]
# *--> Acesse a planilha abaixo (CTRL + click) e insira as informações nas linhas e colunas respectivas*

# %%
# Lance do ano mais recente para o mais antigo
'https://docs.google.com/spreadsheets/d/1qx6nagvF8okAKvfO9BoHL2aZUZ8NOnz2/edit?gid=1575975872#gid=1575975872'

# %% [markdown]
# *--> Agora basta clicar no menu 'Ambiente de execução' e em 'Executar tudo' (ou CTR + F9), nesta ordem.*

# %% [markdown]
# #### 1. CONFIGURAÇÃO DO AMBIENTE

# %% [markdown]
# ##### 1.1 Bibliotecas Python

# %%
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import sys
import subprocess
import importlib.util

from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler

# %%
%%capture
# Lista de pacotes necessários
required_packages = ['gspread', 'pandas', 'gspread_dataframe', 'openpyxl']

# Verificar e instalar pacotes que não estão instalados
def install_missing_packages(packages):
    for package in packages:
        if importlib.util.find_spec(package) is None:
            print(f'⚠ Instalando {package}...')
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', '--upgrade', package])
        else:
            print(f'✔ {package} já está instalado')

install_missing_packages(required_packages)

# %% [markdown]
# ##### 1.2 Fonte dos dados

# %%
# Por ora, uma planilha xlsx no Google Drive (vide 0.4). Futuramente, vincular a um banco de dados, sistema ou API (definir).

# %%
# Definir o ID da planilha e da aba (worksheet)
sheet_id = '1qx6nagvF8okAKvfO9BoHL2aZUZ8NOnz2'
gid = '1575975872'  # ID da aba específica

# Construir a URL para baixar a planilha como um arquivo Excel (.xlsx)
url = f'https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=xlsx&id={sheet_id}&gid={gid}'

# Ler a planilha diretamente no Pandas
df_dados_contabeis = pd.read_excel(url, engine='openpyxl')

# %% [markdown]
# ##### 1.3 Função Índices e Contas Contábeis

# %%
def calcular_indices_contabeis(df):
    indices = {}

    ### RENTABILIDADE
    indices['Margem Líquida'] = df['r_Lucro_Liquido'] / df['r_Receita_Total']
    # Percentual da receita que sobra como lucro líquido.

    indices['ROA'] = df['r_Lucro_Liquido'] / df['p_Ativo_Total']
    # Retorno sobre ativos: eficiência global da operação.

    indices['ROE'] = df['r_Lucro_Liquido'] / df['p_Patrimonio_Liquido']
    # Retorno sobre o patrimônio líquido: rentabilidade para o acionista.

    ### EBITDA
    ebit = df['r_Lucro_Liquido'] + df['r_Despesa_de_Juros'] + df['r_Despesa_de_Impostos']
    amort = df.get('r_Amortizacao', 0).fillna(0)
    depr = df.get('r_Depreciacao', 0).fillna(0)
    ebitda = ebit + amort + depr
    indices['EBITDA'] = ebitda
    indices['Margem EBITDA'] = ebitda / df['r_Receita_Total']
    # Percentual da receita que vira caixa operacional antes de juros, impostos e depreciação.

    ### ALAVANCAGEM E ENDIVIDAMENTO
    df['p_Divida_Bruta'] = df['p_Passivo_Total'] - df['p_Patrimonio_Liquido']
    # Endividamento Bruto (passivo exigível)

    df['p_Divida_Liquida'] = df['p_Divida_Bruta'] - df['p_Caixa']
    # Dívida Líquida (ajustada pelo caixa disponível)

    indices['Alavancagem'] = df['p_Divida_Liquida'] / ebitda
    # Mede quantos anos de geração operacional (EBITDA) seriam necessários para quitar a dívida líquida.

    indices['Endividamento'] = df['p_Divida_Bruta'] / df['p_Ativo_Total']
    # Percentual dos ativos financiado exclusivamente por capital de terceiros.

    ### ESTRUTURA DE CAPITAL
    df['p_Imobilizado'] = df['p_Ativo_Total'] - df['p_Ativo_Circulante']
    # Estimativa do Ativo Imobilizado

    indices['Imobilizado/Ativo'] = df['p_Imobilizado'] / df['p_Ativo_Total']
    # Percentual do ativo total que está imobilizado (capital fixo).

    ### COBERTURA DE JUROS
    indices['Cobertura de Juros'] = ebit / df['r_Despesa_de_Juros']
    # Mede a capacidade da empresa de pagar seus juros com o lucro operacional (EBIT).

    ### EFICIÊNCIA OPERACIONAL
    indices['Giro do Ativo'] = df['r_Receita_Total'] / df['p_Ativo_Total']
    # Mede quantas vezes os ativos se transformam em receita no período.

    ### CICLO OPERACIONAL
    indices['Período Médio de Recebimento'] = df['p_Contas_a_Receber'] / df['r_Receita_Total'] * 365
    indices['Período Médio de Pagamento'] = df['p_Contas_a_Pagar'] / df['r_Custos'] * 365
    # Dias médios para receber e pagar.

    ### LIQUIDEZ
    indices['Liquidez Corrente'] = df['p_Ativo_Circulante'] / df['p_Passivo_Circulante']
    # Mede a capacidade de pagar obrigações de curto prazo.

    indices['Liquidez Seca'] = (df['p_Ativo_Circulante'] - df['p_Estoques']) / df['p_Passivo_Circulante']
    # Medida mais conservadora de liquidez, desconsidera estoques.

    indices['CCL/Ativo Total'] = (df['p_Ativo_Circulante'] - df['p_Passivo_Circulante']) / df['p_Ativo_Total']
    # Mede a folga de capital de giro em relação ao total de ativos da empresa.


    ### Criar DataFrame
    df_indices = pd.DataFrame(indices)

    ### Tratar divisões por zero e valores infinitos
    df_indices.replace([np.inf, -np.inf], np.nan, inplace=True)

    return df_indices


# %% [markdown]
# ##### 1.4 Importação dos Dados Contábeis

# %%
## Importar os dados
# - Do mais recente para o mais antigo

# arquivo_dados_contabeis = '/content/dados_contabeis_global.xlsx'
arquivo_dados_contabeis = df_dados_contabeis

# %%
arquivo_dados_contabeis

# %% [markdown]
# ##### 1.5 Leitura dos Dados Contábeis

# %%
df_dados_contabeis = arquivo_dados_contabeis
print('Dados Contabeis Importados:')
print(df_dados_contabeis)

# %% [markdown]
# #### 2. PROCESSAMENTO DO MODELO #####

# %% [markdown]
# ##### 2.1 Índices contábeis

# %%
# Calculado com base nos dados contábeis

# %%
df_indices = calcular_indices_contabeis(df_dados_contabeis).round(2)
print('\nÍndices Contábeis Calculados:')

# %%
if (df_dados_contabeis['p_Estoques'] == 0).all():
        del df_indices['Liquidez Seca']

# %%
print(df_indices)

# %% [markdown]
# ##### 2.2 Padronização dos índices contábeis

# %%
# Escalar os índices contábeis para o PCA
scaler = StandardScaler()
indices_scaled = scaler.fit_transform(df_indices)
print('\nÍndices Escalados para PCA:')
print(indices_scaled)

# %% [markdown]
# ##### 2.3 Cálculo do PCA

# %%
# Calculado com base nos índices contábeis

# %%
# Realizar o PCA
pca = PCA()
pca_result = pca.fit_transform(indices_scaled)
print('\nComponentes Principais (PCA):')
print(pca_result)

#relatório


# %% [markdown]
# ##### 2.4 Variância Explicada PCA

# %%
# Variância explicada pelos componentes principais
explained_variance_ratio = pca.explained_variance_ratio_
print('\nVariância Explicada por Componente:')
print(explained_variance_ratio)

#relatório

# %% [markdown]
# ##### 2.5 DataFrame PCA

# %%
# DataFrame com os componentes principais
pca_df = pd.DataFrame(pca_result, columns=[f'PC{i+1}' for i in range(pca_result.shape[1])])
print('\nMatriz de Componentes Principais:')
print(pca_df)

#relatório

# %% [markdown]
# ##### 2.6 Matriz de cargas

# %%
# Obter a matriz de cargas
loadings = pd.DataFrame(
    pca.components_,
    columns=df_indices.columns,
    index=[f'PC{i+1}' for i in range(pca.components_.shape[0])]
)

print('Matriz de Cargas dos Componentes Principais:')
print(loadings)

# Identificar os índices mais significativos para cada PC
print('\nÍndices mais significativos por componente:')
for pc in loadings.index:
    print(f'\n{pc}:')
    print(loadings.loc[pc].abs().sort_values(ascending=False).head(3))  # Top 3 índices mais significativos

#relatório

# %%
# Criar dicionário com top 3 índices por componente
top_indices_por_pc = {
    pc: loadings.loc[pc].abs().sort_values(ascending=False).head(3)
    for pc in loadings.index
}

# Transformar em DataFrame
top_indices_df = pd.DataFrame([
    {
        'PC': pc,
        'Indice 1': top.index[0],
        'Peso 1': round(top.values[0], 3),
        'Indice 2': top.index[1],
        'Peso 2': round(top.values[1], 3),
        'Indice 3': top.index[2],
        'Peso 3': round(top.values[2], 3)
    }
    for pc, top in top_indices_por_pc.items()
])

# Exibir ou salvar
print('\nTop 3 Índices por Componente Principal:')
print(top_indices_df)


# %% [markdown]
# ##### 2.7 Score PCA final

# %%
pca_df.dot(explained_variance_ratio)

# %% [markdown]
# #### 3. RESULTADOS #####

# %% [markdown]
# 

# %% [markdown]
# ##### 3.1A Função Finscore Bruto

# %%
# Função para categorizar escores consolidados com mais granularidade
def categorias_finscore_bruto(escores):
    categorias = []
    for escore in escores:
        if escore > 1.5:
            categorias.append('Muito Abaixo do Risco')
        elif 1.0 < escore <= 1.5:
            categorias.append('Levemente Abaixo do Risco')
        elif -1.0 <= escore <= 1.0:
            categorias.append('Neutro')
        elif -1.5 < escore < -1.0:
            categorias.append('Levemente Acima do Risco')
        else:
            categorias.append('Muito Acima do Risco')
    return categorias

# Tabela de Categorias:
# -------------------------------------------------------
# |   Intervalo do Escore   |      Categoria            |
# -------------------------------------------------------
# |  escore > 1.5           | Muito Abaixo do Risco     |
# |  1.0 < escore ≤ 1.5     | Levemente Abaixo do Risco |
# | -1.0 ≤ escore ≤ 1.0     | Neutro                    |
# | -1.5 < escore < -1.0    | Levemente Acima do Risco  |
# |  escore ≤ -1.5          | Muito Acima do Risco      |
# -------------------------------------------------------

# %% [markdown]
# ##### 3.1B Calcular Finscore Bruto

# %%
# Calcular o escore consolidado com penalização do último ano
# Do mais recente para o mais antigo: 
pesos = [0.6, 0.25, 0.15]  # Pesos para os três anos
finscore_bruto = round((pca_df.dot(explained_variance_ratio) * pesos).sum(), 2)
print('\nFINSCORE BRUTO:')
print(finscore_bruto)
print('\nCLASSIFICAÇÃO FINSCORE BRUTO:')
print(categorias_finscore_bruto([finscore_bruto])[0])


# %% [markdown]
# ##### 3.2A Função Finscore Ajustado

# %%
# Este é o finscore que será considerado para o relatório.

# %%
# Função para categorizar escores consolidados com mais granularidade
def categorias_finscore_ajustado(escore):
    if escore > 875:
        return 'Muito Abaixo do Risco'
    elif 750 < escore <= 875:
        return 'Levemente Abaixo do Risco'
    elif 250 <= escore <= 750:
        return 'Neutro'
    elif 125 < escore < 250:
        return 'Levemente Acima do Risco'
    else:
        return 'Muito Acima do Risco'


# Tabela de Categorias:
# -------------------------------------------------------
# |   Intervalo do Escore  |      Categoria            |
# -------------------------------------------------------
# |  escore > 750          | Muito Abaixo do Risco     |
# |  500 < escore ≤ 750    | Levemente Abaixo do Risco |
# | 250 ≤ escore ≤ 500     | Neutro                    |
# | 100 < escore < 250     | Levemente Acima do Risco  |
# |  escore ≤ 100          | Muito Acima do Risco      |
# -------------------------------------------------------

# %% [markdown]
# ##### 3.2B Calcular Finscore Ajustado

# %%
# Traduz a tendência em um escore de risco ajustado ao ponto de referência atual do cliente.
finscore_ajustado = round(min(((finscore_bruto + 2)/4)*1000, 1000), 2)

# Converte escore bruto (-2 a +2) para uma escala de 0 a 1000, truncando em 1000.


# %%
print('\nFINSCORE AJUSTADO:')
print(finscore_ajustado)
print('\nCLASSIFICAÇÃO FINSCORE AJUSTADO:')
print(categorias_finscore_ajustado(finscore_ajustado))

# %% [markdown]
# ##### 3.3A Função Serasa

# %%
# Função para categorizar escores conforme a classificação do Serasa
def categorias_serasa(score):
    if score > 700:
        return 'Excelente'
    elif score > 500:
        return 'Bom'
    elif score > 300:
        return 'Baixo'
    else:
        return 'Muito Baixo'

# Tabela de Categorias:
# --------------------------------------
# |   Intervalo do Escore  | Categoria |
# --------------------------------------
# |  701 a 1.000          | Excelente  |
# |  501 a 700            | Bom        |
# |  301 a 500            | Baixo      |
# |  0 a 300              | Muito Baixo|
# --------------------------------------

# %% [markdown]
# ##### 3.3B Calcular Serasa

# %%
print('\nSERASA:')
print(serasa)
print('\nCLASSIFICAÇÃO SERASA:')
print(categorias_serasa(serasa))

# %% [markdown]
# ##### 3.4 Valores e Contas Contábeis

# %%
# Apenas para revisão

# %%
df_dados_contabeis
# Sendo 0 o mais recente e 2 o mais antigo

# %% [markdown]
# ##### 3.5 Índices contábeis

# %%
# Apenas para revisão

# %%
print(df_indices)
# Sendo 0 o mais recente e 2 o mais antigo


# %% [markdown]
# #### 4. PREPARAÇÃO IA

# %%
# Prompts e funções para o ChatGPT e LangChain

# %% [markdown]
# ##### 4.1 Dicionário IA

# %%
# Criar o DataFrame com os resultados principais para IA
resultados_df_final = pd.DataFrame({
    'Métrica': ['Finscore', 'Serasa'],
    'Valor': [finscore_ajustado, serasa],
    'Categoria': [
        categorias_finscore_ajustado(finscore_ajustado),
        categorias_serasa(serasa)
    ]
})


# %%
# Função para preparar os dados em formato de dicionário para a IA
def preparar_para_llm(df_indices, resultados_df_final, top_indices_df, df_dados_contabeis):
    return {
        'cliente': cliente,
        'periodo_analise': f'{ano_inicial}–{ano_final}',
        'indices_financeiros': df_indices.to_dict(orient='records'),
        'finScore_resultado': resultados_df_final.to_dict(orient='records'),
        'pca_destaques': top_indices_df.to_dict(orient='records'),
        'dados_contabeis': df_dados_contabeis.to_dict(orient='records')
    }


# %% [markdown]
# ##### 4.2 Chamar dados IA

# %%
# Chamada da função com os dados disponíveis
contexto_ia = preparar_para_llm(
    df_indices=df_indices,
    resultados_df_final=resultados_df_final,
    top_indices_df=top_indices_df,
    df_dados_contabeis=df_dados_contabeis
)

# %% [markdown]
# ##### 4.3 Visualizar Dados IA

# %%
# Visualização opcional do dicionário gerado (útil para debug ou checagem)
import json
print('\n📦 DADOS PREPARADOS PARA IA (formato .json simplificado):')
print(json.dumps(contexto_ia, indent=2, ensure_ascii=False)[:3000])  # mostra só os primeiros 3000 caracteres

# %%
# A versão com llm.invoke é ótima para validar o texto direto.

# A versão com LLMChain vai te preparar para expandir depois (ex: exportar para Word, gerar resumo técnico, enviar via API.

# %% [markdown]
# #### 5. PARECER IA

# %%
# https://platform.openai.com/docs/overview
# https://platform.openai.com/settings/organization/billing/history
# https://platform.openai.com/settings/organization/usage

# %% [markdown]
# ##### 5.1 Promp Finscore

# %%
from langchain.prompts import PromptTemplate

prompt_template = """
Você é um analista de crédito sênior de um banco de desenvolvimento. Recebeu os dados financeiros consolidados de uma empresa, contendo contas contábeis, índices financeiros, componentes principais (PCA) e os escores de risco FinScore e Serasa.

Com base nesses dados, elabore um **parecer técnico completo** com os seguintes critérios:

1. **INTRODUÇÃO**
- Apresente o objetivo do parecer.
- Identifique o cliente, o período da análise e o escopo da avaliação (contábil, econômico-financeira e de risco).

2. **METODOLOGIA**
- Explique o que é o FinScore, como é calculado (PCA + índices contábeis) e sua utilidade para análise de crédito.
- Explique o Serasa Score como índice complementar e relacione ambos no contexto do diagnóstico de risco.
- Diga que os índices foram padronizados e ponderados no tempo, destacando os principais vetores do PCA.

3. **ANÁLISE DAS CONTAS CONTÁBEIS**
- Apresente uma **tabela com os valores principais** (ativos, passivos, lucro, receitas, caixa etc.).
- Comente **criticamente a evolução** desses valores ao longo do tempo, apontando **pontos fortes, riscos ou deteriorações**.

4. **ANÁLISE DOS ÍNDICES FINANCEIROS**
- Mostre uma **tabela com todos os índices calculados**.
- Comente um a um: liquidez, rentabilidade, alavancagem, giro, cobertura de juros, prazos médios etc.
- Destaque **os três mais relevantes** segundo o PCA e interprete **por que eles se destacam** e o que indicam.

5. **FINSCORE E SERASA SCORE**
- Apresente os dois scores obtidos, em formato **tabela com classificação**.
- Comente **as classificações e coerência entre os escores**.
- Relacione os escores com os dados contábeis e com a saúde financeira geral da empresa.

6. **CONCLUSÃO E VEREDICTO FINAL**
- Diga se recomenda a concessão de crédito (sim, não, sim com ressalvas).
- Fundamente a decisão com base em dados objetivos e análises anteriores.
- Seja **criterioso e técnico**, considerando o risco vs. retorno para a instituição.

7. **RECOMENDAÇÕES PARA MITIGAÇÃO DE RISCOS**
- Sugira garantias, limites prudenciais, cláusulas contratuais, acompanhamento contínuo do FinScore e envio trimestral de dados financeiros.
- Se necessário, proponha exigência de avalistas ou reforço de capital.

**IMPORTANTE:**  
- Escreva com clareza, formalidade e linguagem técnica.
- Insira as **tabelas e gráficos onde for pertinente** (mesmo que simbólicos, como: [Tabela: Índices Financeiros], [Gráfico: Evolução do Endividamento]).
- Evite redundâncias. Destaque **o que os números revelam**.
- Apresente **insights concretos** para o tomador de decisão.

DADOS DISPONÍVEIS:
- Cliente: {cliente}
- Período da análise: {periodo_analise}
- Contas contábeis: {dados_contabeis}
- Índices financeiros: {indices_financeiros}
- FinScore e Serasa: {finScore_resultado}
- PCA (destaques): {pca_destaques}
"""

# Criar o PromptTemplate
prompt = PromptTemplate.from_template(prompt_template)


# %% [markdown]
# ##### 5.2 Chave OpenAI

# %%
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import os

# Carregar variáveis do .env
load_dotenv()
api_key = os.getenv('OPENAI_API_KEY')
print(f'API_KEY carregada: {api_key is not None}')

# Verificar se a chave foi carregada corretamente
if api_key:
    print('OPENAI_API_KEY carregada com sucesso.')
else:
    print('OPENAI_API_KEY não encontrada. Verifique seu arquivo .env.')

# %%
print("Diretório atual:", os.getcwd())

# %% [markdown]
# ##### 5.3 Configurar modelo OpenAI com LangChain

# %%
from langchain_community.chat_models import ChatOpenAI

# Escolher o modelo a ser utilizado
modelo = 4 # '3.5' ou '4'

# ifelse modelo gpt-3.5 ou gpt-4
if modelo == 3.5:
    model = 'gpt-3.5-turbo'
elif modelo == 4:
    model = 'gpt-4'
else:
    raise ValueError('Modelo inválido. Escolha entre 3.5 ou 4.')

# Cria uma instância do modelo ChatGPT da OpenAI
llm = ChatOpenAI(
    temperature=0.2,       # Baixa temperatura → respostas mais objetivas e determinísticas
    model=model,           # Usando o modelo definido acima
    openai_api_key=api_key
)

print('Modelo ChatOpenAI configurado com sucesso.')


# %% [markdown]
# ##### 5.4 Gerar o relatório

# %%
prompt_final = prompt.format(**contexto_ia)
resposta = llm.invoke(prompt_final)
print(resposta.content) 


# %% [markdown]
# #### 6. Exportação Word

# %% [markdown]
# ##### 6.1 Importar Bibliotecas

# %%
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# %% [markdown]
# ##### 6.2 Criar o documento

# %%
# Remover título duplicado do conteúdo, se houver
resposta.content = re.sub(r'^\*?\*?Parecer Técnico de Análise de Crédito\*?\*?\s*', '', resposta.content.strip(), flags=re.IGNORECASE)


# %%
import re
from docx.enum.style import WD_STYLE_TYPE

# Criar o documento
doc = Document()
doc.add_heading('Parecer Técnico de Análise de Crédito', level=1)

# Criar estilo de parágrafo justificado
style = doc.styles.add_style('Justificado', WD_STYLE_TYPE.PARAGRAPH)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Dividir e formatar conteúdo
secoes = resposta.content.split('\n\n')

for secao in secoes:
    texto = secao.strip()
    if not texto:
        continue

    # Detectar título de seção
    match = re.match(r'^\*\*(\d\..*?)\*\*$', texto)
    if match:
        doc.add_heading(match.group(1), level=2)
    else:
        paragrafo = doc.add_paragraph(texto)
        paragrafo.style = 'Justificado'
        for run in paragrafo.runs:
            run.font.size = Pt(11)


# %% [markdown]
# ##### 6.3 Gerar o documento

# %%
# Salvar o arquivo
caminho_arquivo = f'Parecer_FinScore_{cliente.replace(" ", "_")}_{ano_final}.docx'
doc.save(caminho_arquivo)

print(f' Documento exportado com sucesso: {caminho_arquivo}')


